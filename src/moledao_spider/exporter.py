"""DOCX exporter that batches jobs into groups of ten."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Iterable, List, Sequence

from bs4 import BeautifulSoup
from docx import Document

from .models import CareerRecord

logger = logging.getLogger(__name__)


def _chunk(items: Sequence[CareerRecord], size: int) -> Iterable[Sequence[CareerRecord]]:
    for idx in range(0, len(items), size):
        yield items[idx : idx + size]


def _html_to_lines(html: str) -> List[str]:
    if not html:
        return ["N/A"]
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()

    lines: List[str] = []
    for block in soup.find_all(["p", "li", "div"]):
        text = block.get_text(" ", strip=True)
        if text:
            lines.append(text)

    if not lines:
        text = soup.get_text(" ", strip=True)
        if text:
            lines.append(text)

    return lines or ["N/A"]


class DocxExporter:
    def __init__(self, output_dir: Path, batch_size: int = 10, append: bool = False):
        self.output_dir = output_dir
        self.batch_size = max(1, batch_size)
        self.append = append

    def export(self, jobs: Sequence[CareerRecord]) -> List[Path]:
        if not jobs:
            logger.warning("No jobs provided to exporter")
            return []
        self.output_dir.mkdir(parents=True, exist_ok=True)
        written: List[Path] = []

        start_index = self._next_index() if self.append else 1

        for offset, chunk in enumerate(_chunk(jobs, self.batch_size)):
            document = Document()
            for idx, job in enumerate(chunk):
                self._write_job(document, job)
                if idx < len(chunk) - 1:
                    document.add_paragraph()
            filename = self._filename(start_index + offset)
            if not self.append and filename.exists():
                logger.info("Overwriting %s", filename)
            document.save(filename)
            written.append(filename)
            logger.info("Wrote %s", filename)

        return written

    def _filename(self, doc_index: int) -> Path:
        return self.output_dir / f"jobs-{doc_index:03d}.docx"

    def _next_index(self) -> int:
        existing = sorted(self.output_dir.glob("jobs-*.docx"))
        pattern = re.compile(r"jobs-(\d{3})\.docx$")
        max_index = 0
        for path in existing:
            match = pattern.match(path.name)
            if match:
                max_index = max(max_index, int(match.group(1)))
        return max_index + 1 if max_index else 1

    def _write_job(self, document: Document, job: CareerRecord) -> None:
        document.add_heading(job.company, level=1)
        document.add_heading(f"{job.role} ({job.type_text})", level=2)
        document.add_paragraph(f"Location: {job.location}")
        document.add_paragraph(f"Type: {job.type_text}")
        document.add_paragraph(f"Preferences: {job.preference_text}")
        document.add_paragraph(job.relative_time)
        document.add_paragraph(f"Exp: {job.experience_text}")
        document.add_paragraph(f"Tag: {job.tag_text}")
        document.add_paragraph("content:")
        for line in _html_to_lines(job.html_content):
            document.add_paragraph(line)
        document.add_paragraph(f"time: {job.update_date}")
